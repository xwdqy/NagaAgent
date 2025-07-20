"""
Formatting tools for Word Document Server.

These tools handle formatting operations for Word documents,
including text formatting, table formatting, and custom styles.
"""
import os
from typing import List, Optional, Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
from word_document_server.utils.document_utils import find_and_replace_text
from word_document_server.core.styles import create_style
from word_document_server.core.tables import apply_table_style


async def format_text(filename: str, paragraph_index: int, start_pos: int, end_pos: int, 
                     bold: Optional[bool] = None, italic: Optional[bool] = None, 
                     underline: Optional[bool] = None, color: Optional[str] = None,
                     font_size: Optional[int] = None, font_name: Optional[str] = None) -> str:
    """Format a specific range of text within a paragraph.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph (0-based)
        start_pos: Start position within the paragraph text
        end_pos: End position within the paragraph text
        bold: Set text bold (True/False)
        italic: Set text italic (True/False)
        underline: Set text underlined (True/False)
        color: Text color (e.g., 'red', 'blue', etc.)
        font_size: Font size in points
        font_name: Font name/family
    """
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        paragraph_index = int(paragraph_index)
        start_pos = int(start_pos)
        end_pos = int(end_pos)
        if font_size is not None:
            font_size = int(font_size)
    except (ValueError, TypeError):
        return "Invalid parameter: paragraph_index, start_pos, end_pos, and font_size must be integers"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        paragraph = doc.paragraphs[paragraph_index]
        text = paragraph.text
        
        # Validate text positions
        if start_pos < 0 or end_pos > len(text) or start_pos >= end_pos:
            return f"Invalid text positions. Paragraph has {len(text)} characters."
        
        # Get the text to format
        target_text = text[start_pos:end_pos]
        
        # Clear existing runs and create three runs: before, target, after
        for run in paragraph.runs:
            run.clear()
        
        # Add text before target
        if start_pos > 0:
            run_before = paragraph.add_run(text[:start_pos])
        
        # Add target text with formatting
        run_target = paragraph.add_run(target_text)
        if bold is not None:
            run_target.bold = bold
        if italic is not None:
            run_target.italic = italic
        if underline is not None:
            run_target.underline = underline
        if color:
            # Define common RGB colors
            color_map = {
                'red': RGBColor(255, 0, 0),
                'blue': RGBColor(0, 0, 255),
                'green': RGBColor(0, 128, 0),
                'yellow': RGBColor(255, 255, 0),
                'black': RGBColor(0, 0, 0),
                'gray': RGBColor(128, 128, 128),
                'white': RGBColor(255, 255, 255),
                'purple': RGBColor(128, 0, 128),
                'orange': RGBColor(255, 165, 0)
            }
            
            try:
                if color.lower() in color_map:
                    # Use predefined RGB color
                    run_target.font.color.rgb = color_map[color.lower()]
                else:
                    # Try to set color by name
                    run_target.font.color.rgb = RGBColor.from_string(color)
            except Exception as e:
                # If all else fails, default to black
                run_target.font.color.rgb = RGBColor(0, 0, 0)
        if font_size:
            run_target.font.size = Pt(font_size)
        if font_name:
            run_target.font.name = font_name
        
        # Add text after target
        if end_pos < len(text):
            run_after = paragraph.add_run(text[end_pos:])
        
        doc.save(filename)
        return f"Text '{target_text}' formatted successfully in paragraph {paragraph_index}."
    except Exception as e:
        return f"Failed to format text: {str(e)}"


async def create_custom_style(filename: str, style_name: str, 
                             bold: Optional[bool] = None, italic: Optional[bool] = None,
                             font_size: Optional[int] = None, font_name: Optional[str] = None,
                             color: Optional[str] = None, base_style: Optional[str] = None) -> str:
    """Create a custom style in the document.
    
    Args:
        filename: Path to the Word document
        style_name: Name for the new style
        bold: Set text bold (True/False)
        italic: Set text italic (True/False)
        font_size: Font size in points
        font_name: Font name/family
        color: Text color (e.g., 'red', 'blue')
        base_style: Optional existing style to base this on
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Build font properties dictionary
        font_properties = {}
        if bold is not None:
            font_properties['bold'] = bold
        if italic is not None:
            font_properties['italic'] = italic
        if font_size is not None:
            font_properties['size'] = font_size
        if font_name is not None:
            font_properties['name'] = font_name
        if color is not None:
            font_properties['color'] = color
        
        # Create the style
        new_style = create_style(
            doc, 
            style_name, 
            WD_STYLE_TYPE.PARAGRAPH, 
            base_style=base_style,
            font_properties=font_properties
        )
        
        doc.save(filename)
        return f"Style '{style_name}' created successfully."
    except Exception as e:
        return f"Failed to create style: {str(e)}"


async def format_table(filename: str, table_index: int, 
                      has_header_row: Optional[bool] = None,
                      border_style: Optional[str] = None,
                      shading: Optional[List[List[str]]] = None) -> str:
    """Format a table with borders, shading, and structure.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        has_header_row: If True, formats the first row as a header
        border_style: Style for borders ('none', 'single', 'double', 'thick')
        shading: 2D list of cell background colors (by row and column)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Apply formatting
        success = apply_table_style(table, has_header_row, border_style, shading)
        
        if success:
            doc.save(filename)
            return f"Table at index {table_index} formatted successfully."
        else:
            return f"Failed to format table at index {table_index}."
    except Exception as e:
        return f"Failed to format table: {str(e)}"
