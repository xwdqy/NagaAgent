import json
import os
from typing import Dict, Any, Optional

# 简化的Word文档处理功能
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx未安装，Word文档功能将受限")

class WordDocumentMCPServer:
    """Word文档处理MCP服务器适配器"""
    
    def __init__(self):
        self.name = "WordDocumentMCPServer"
        self.instructions = "专业的Microsoft Word文档创建、编辑和管理工具"
        
        # 工具映射表
        self.tool_mapping = {
            # 文档管理
            "create_document": self._create_document,
            "get_document_info": self._get_document_info,
            "get_document_text": self._get_document_text,
            "list_available_documents": self._list_available_documents,
            
            # 内容添加
            "add_paragraph": self._add_paragraph,
            "add_heading": self._add_heading,
            "add_table": self._add_table,
            "add_page_break": self._add_page_break,
            
            # 其他功能
            "help": self._help
        }
    
    async def handle_handoff(self, data: dict) -> str:
        """处理MCP工具调用"""
        try:
            tool_name = data.get("tool_name")
            if not tool_name:
                return self._error_response("缺少tool_name参数")
            
            if tool_name not in self.tool_mapping:
                return self._error_response(f"未知工具: {tool_name}，可用工具: {list(self.tool_mapping.keys())}")
            
            # 调用对应的工具函数
            handler = self.tool_mapping[tool_name]
            result = handler(data)
            
            return json.dumps({
                "status": "ok",
                "message": "操作成功",
                "data": result
            }, ensure_ascii=False)
            
        except Exception as e:
            return self._error_response(f"操作失败: {str(e)}")
    
    def _error_response(self, message: str) -> str:
        """生成错误响应"""
        return json.dumps({
            "status": "error",
            "message": message,
            "data": ""
        }, ensure_ascii=False)
    
    def _help(self, data: dict) -> dict:
        """显示帮助信息"""
        return {
            "available_tools": list(self.tool_mapping.keys()),
            "description": "Word文档处理MCP服务器",
            "docx_available": DOCX_AVAILABLE,
            "examples": [
                {
                    "tool": "create_document",
                    "description": "创建新文档",
                    "params": {"filename": "test.docx", "title": "测试文档"}
                },
                {
                    "tool": "add_paragraph",
                    "description": "添加段落",
                    "params": {"filename": "test.docx", "text": "这是一个段落"}
                }
            ]
        }
    
    def _create_document(self, data: dict) -> dict:
        """创建新的Word文档，支持自定义保存路径"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx未安装，无法创建Word文档")
        
        # 支持多种参数名：file_path, filename
        filename = data.get("file_path") or data.get("filename", "document.docx")
        title = data.get("title", "")
        author = data.get("author", "")
        save_path = data.get("save_path", None)  # 新增保存路径参数
        
        try:
            # 确保文件名以.docx结尾
            if not filename.endswith('.docx'):
                filename += '.docx'
            
            # 解析完整路径
            if save_path:
                # 规范化路径并确保目录存在
                save_path = os.path.abspath(save_path)
                os.makedirs(save_path, exist_ok=True)
                full_path = os.path.join(save_path, filename)
            else:
                full_path = filename
            
            # 创建新文档
            doc = Document()
            
            # 设置文档属性
            if title:
                doc.core_properties.title = title
            if author:
                doc.core_properties.author = author
            
            # 添加文档标题（如果提供了title）
            if title:
                doc.add_heading(title, level=1)
            
            # 如果提供了content参数，添加内容段落
            content = data.get("content", "")
            if content:
                doc.add_paragraph(content)
            
            # 如果作者信息存在，添加作者段落
            if author:
                doc.add_paragraph(f'Author: {author}')
            
            # 保存文档
            doc.save(full_path)
            
            return {
                "filename": os.path.basename(full_path),
                "directory": os.path.dirname(full_path) if save_path else ".",
                "full_path": os.path.abspath(full_path),
                "title": title,
                "author": author,
                "created": True,
                "save_path": save_path
            }
        except Exception as e:
            raise Exception(f"创建文档失败: {str(e)}")
    
    def _add_paragraph(self, data: dict) -> dict:
        """添加段落"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx未安装，无法编辑Word文档")
        
        # 支持多种参数名：file_path, filename
        filename = data.get("file_path") or data.get("filename")
        text = data.get("text", "")
        
        if not filename:
            raise Exception("缺少file_path或filename参数")
        
        try:
            # 打开文档
            if os.path.exists(filename):
                doc = Document(filename)
            else:
                raise FileNotFoundError(f"文档 '{filename}' 不存在")
            
            # 添加段落
            paragraph = doc.add_paragraph(text)
            
            # 保存文档
            doc.save(filename)
            
            return {
                "filename": filename,
                "text": text,
                "added": True
            }
        except Exception as e:
            raise Exception(f"添加段落失败: {str(e)}")
    
    def _add_heading(self, data: dict) -> dict:
        """添加标题"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx未安装，无法编辑Word文档")
        
        # 支持多种参数名：file_path, filename
        filename = data.get("file_path") or data.get("filename")
        text = data.get("text", "")
        level = data.get("level", 1)
        
        if not filename:
            raise Exception("缺少file_path或filename参数")
        
        # 确保level是整数类型
        try:
            level = int(level)
        except (ValueError, TypeError):
            raise Exception("level参数必须是1-9之间的整数")
        
        # 验证level范围
        if level < 1 or level > 9:
            raise Exception(f"无效的标题级别: {level}。级别必须在1-9之间。")
        
        try:
            # 打开文档
            if os.path.exists(filename):
                doc = Document(filename)
            else:
                raise FileNotFoundError(f"文档 '{filename}' 不存在")
            
            # 添加标题
            heading = doc.add_heading(text, level)
            
            # 保存文档
            doc.save(filename)
            
            return {
                "filename": filename,
                "text": text,
                "level": level,
                "added": True
            }
        except Exception as e:
            raise Exception(f"添加标题失败: {str(e)}")
    
    def _add_table(self, data: dict) -> dict:
        """添加表格"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx未安装，无法编辑Word文档")
        
        # 支持多种参数名：file_path, filename
        filename = data.get("file_path") or data.get("filename")
        rows = data.get("rows", 2)
        cols = data.get("cols", 2)
        headers = data.get("headers", [])
        
        if not filename:
            raise Exception("缺少file_path或filename参数")
        
        # 确保rows和cols是整数类型
        try:
            rows = int(rows)
            cols = int(cols)
        except (ValueError, TypeError):
            raise Exception("rows和cols参数必须是正整数")
        
        # 验证rows和cols范围
        if rows < 1 or cols < 1:
            raise Exception(f"无效的表格尺寸: {rows}x{cols}。行数和列数必须大于0。")
        
        if rows > 100 or cols > 50:  # 设置合理的上限
            raise Exception(f"表格尺寸过大: {rows}x{cols}。建议行数不超过100，列数不超过50。")
        
        try:
            # 打开文档
            if os.path.exists(filename):
                doc = Document(filename)
            else:
                raise FileNotFoundError(f"文档 '{filename}' 不存在")
            
            # 添加表格
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            # 如果提供了表头，设置第一行
            if headers and len(headers) <= cols:
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    if i < len(hdr_cells):
                        hdr_cells[i].text = str(header)
            
            # 保存文档
            doc.save(filename)
            
            return {
                "filename": filename,
                "rows": rows,
                "cols": cols,
                "headers": headers,
                "added": True
            }
        except Exception as e:
            raise Exception(f"添加表格失败: {str(e)}")
    
    def _add_page_break(self, data: dict) -> dict:
        """添加分页符"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx未安装，无法编辑Word文档")
        
        # 支持多种参数名：file_path, filename
        filename = data.get("file_path") or data.get("filename")
        
        if not filename:
            raise Exception("缺少file_path或filename参数")
        
        try:
            # 打开文档
            if os.path.exists(filename):
                doc = Document(filename)
            else:
                raise FileNotFoundError(f"文档 '{filename}' 不存在")
            
            # 添加分页符
            doc.add_page_break()
            
            # 保存文档
            doc.save(filename)
            
            return {
                "filename": filename,
                "added": True
            }
        except Exception as e:
            raise Exception(f"添加分页符失败: {str(e)}")
    
    def _get_document_info(self, data: dict) -> dict:
        """获取文档信息"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx未安装，无法读取Word文档")
        
        # 支持多种参数名：file_path, filename
        filename = data.get("file_path") or data.get("filename")
        
        if not filename:
            raise Exception("缺少file_path或filename参数")
        
        try:
            if not os.path.exists(filename):
                raise FileNotFoundError(f"文档 '{filename}' 不存在")
            
            # 打开文档
            doc = Document(filename)
            
            # 获取文档属性
            props = doc.core_properties
            
            return {
                "filename": filename,
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "paragraph_count": len(doc.paragraphs),
                "file_size": os.path.getsize(filename)
            }
        except Exception as e:
            raise Exception(f"获取文档信息失败: {str(e)}")
    
    def _get_document_text(self, data: dict) -> dict:
        """获取文档文本内容"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx未安装，无法读取Word文档")
        
        # 支持多种参数名：file_path, filename
        filename = data.get("file_path") or data.get("filename")
        
        if not filename:
            raise Exception("缺少file_path或filename参数")
        
        try:
            if not os.path.exists(filename):
                raise FileNotFoundError(f"文档 '{filename}' 不存在")
            
            # 打开文档
            doc = Document(filename)
            
            # 提取所有段落文本
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # 只添加非空段落
                    text_content.append(paragraph.text)
            
            return {
                "filename": filename,
                "text": "\n".join(text_content),
                "paragraph_count": len(text_content)
            }
        except Exception as e:
            raise Exception(f"获取文档文本失败: {str(e)}")
    
    def _list_available_documents(self, data: dict) -> dict:
        """列出可用文档"""
        directory = data.get("directory", ".")
        
        try:
            if not os.path.exists(directory):
                raise FileNotFoundError(f"目录 '{directory}' 不存在")
            
            # 查找所有.docx文件
            docx_files = []
            for file in os.listdir(directory):
                if file.endswith('.docx') and not file.startswith('~$'):
                    file_path = os.path.join(directory, file)
                    file_info = {
                        "filename": file,
                        "path": file_path,
                        "size": os.path.getsize(file_path),
                        "modified": os.path.getmtime(file_path)
                    }
                    docx_files.append(file_info)
            
            return {
                "directory": directory,
                "documents": docx_files,
                "count": len(docx_files)
            }
        except Exception as e:
            raise Exception(f"列出文档失败: {str(e)}")

def create_word_document_mcp_server():
    """创建Word文档MCP服务器实例"""
    return WordDocumentMCPServer()