"""
Table-related operations for Word Document Server.
"""
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def set_cell_border(cell, **kwargs):
    """
    Set cell border properties.
    
    Args:
        cell: The cell to modify
        **kwargs: Border properties (top, bottom, left, right, val, color)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create border elements
    for key, value in kwargs.items():
        if key in ['top', 'left', 'bottom', 'right']:
            tag = 'w:{}'.format(key)
            
            element = OxmlElement(tag)
            element.set(qn('w:val'), kwargs.get('val', 'single'))
            element.set(qn('w:sz'), kwargs.get('sz', '4'))
            element.set(qn('w:space'), kwargs.get('space', '0'))
            element.set(qn('w:color'), kwargs.get('color', 'auto'))
            
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
                
            tcBorders.append(element)


def apply_table_style(table, has_header_row=False, border_style=None, shading=None):
    """
    Apply formatting to a table.
    
    Args:
        table: The table to format
        has_header_row: If True, formats the first row as a header
        border_style: Style for borders ('none', 'single', 'double', 'thick')
        shading: 2D list of cell background colors (by row and column)
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Format header row if requested
        if has_header_row and table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.runs:
                        for run in paragraph.runs:
                            run.bold = True
        
        # Apply border style if specified
        if border_style:
            val_map = {
                'none': 'nil',
                'single': 'single',
                'double': 'double',
                'thick': 'thick'
            }
            val = val_map.get(border_style.lower(), 'single')
            
            # Apply to all cells
            for row in table.rows:
                for cell in row.cells:
                    set_cell_border(
                        cell,
                        top=True,
                        bottom=True,
                        left=True,
                        right=True,
                        val=val,
                        color="000000"
                    )
        
        # Apply cell shading if specified
        if shading:
            for i, row_colors in enumerate(shading):
                if i >= len(table.rows):
                    break
                for j, color in enumerate(row_colors):
                    if j >= len(table.rows[i].cells):
                        break
                    try:
                        # Apply shading to cell
                        cell = table.rows[i].cells[j]
                        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
                        cell._tc.get_or_add_tcPr().append(shading_elm)
                    except:
                        # Skip if color format is invalid
                        pass
        
        return True
    except Exception:
        return False


def copy_table(source_table, target_doc):
    """
    Copy a table from one document to another.
    
    Args:
        source_table: The table to copy
        target_doc: The document to copy the table to
        
    Returns:
        The new table in the target document
    """
    # Create a new table with the same dimensions
    new_table = target_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
    
    # Try to apply the same style
    try:
        if source_table.style:
            new_table.style = source_table.style
    except:
        # Fall back to default grid style
        try:
            new_table.style = 'Table Grid'
        except:
            pass
    
    # Copy cell contents
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                if paragraph.text:
                    new_table.cell(i, j).text = paragraph.text
    
    return new_table
