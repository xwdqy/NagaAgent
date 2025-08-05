"""
Footnote and endnote functionality for Word Document Server.
"""
from docx import Document
from typing import List, Tuple


def add_footnote(doc, paragraph, text):
    """
    Add a footnote to a paragraph.
    
    Args:
        doc: Document object
        paragraph: Paragraph to add footnote to
        text: Text content of the footnote
    
    Returns:
        The created footnote
    """
    return paragraph.add_footnote(text)


def add_endnote(doc, paragraph, text):
    """
    Add an endnote to a paragraph.
    This is a custom implementation since python-docx doesn't directly support endnotes.
    
    Args:
        doc: Document object
        paragraph: Paragraph to add endnote to
        text: Text content of the endnote
    
    Returns:
        The paragraph containing the endnote reference
    """
  
    run = paragraph.add_run()
    run.text = "*"
    run.font.superscript = True
    
    # Add endnote text at the end of the document
    # create a section for endnotes if it doesn't exist
    endnotes_found = False
    for para in doc.paragraphs:
        if para.text == "Endnotes:":
            endnotes_found = True
            break
    
    if not endnotes_found:
        # Add a page break before endnotes section
        doc.add_page_break()
        doc.add_heading("Endnotes:", level=1)
    
    # Add the endnote text
    endnote_text = f"* {text}"
    doc.add_paragraph(endnote_text)
    
    return paragraph


def convert_footnotes_to_endnotes(doc):
    """
    Convert all footnotes to endnotes in a document.
    
    Args:
        doc: Document object
    
    Returns:
        Number of footnotes converted
    """
    # This is a complex operation not fully supported by python-docx
    # Implementing a simplified version
    
    # Collect all footnotes
    footnotes = []
    for para in doc.paragraphs:
        
        # This is a simplified implementation
        for run in para.runs:
            if run.font.superscript and run.text.isdigit():
                # This might be a footnote reference
                footnotes.append((para, run.text))
    
    # Add endnotes section
    if footnotes:
        doc.add_page_break()
        doc.add_heading("Endnotes:", level=1)
        
        # Add each footnote as an endnote
        for idx, (para, footnote_num) in enumerate(footnotes):
            doc.add_paragraph(f"{idx+1}. Converted from footnote {footnote_num}")
    
    return len(footnotes)


def find_footnote_references(doc) -> List[Tuple[int, int, str]]:
    """
    Find all footnote references in a document.
    
    Args:
        doc: Document object
        
    Returns:
        List of tuples (paragraph_index, run_index, text) for each footnote reference
    """
    footnote_references = []
    
    for para_idx, para in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(para.runs):
           
            if run.font.superscript and (run.text.isdigit() or run.text in "¹²³⁴⁵⁶⁷⁸⁹"):
                footnote_references.append((para_idx, run_idx, run.text))
    
    return footnote_references


def get_format_symbols(numbering_format: str, count: int) -> List[str]:
    """
    Get a list of formatting symbols based on the specified numbering format.
    
    Args:
        numbering_format: Format for footnote/endnote numbers (e.g., "1, 2, 3", "i, ii, iii", "a, b, c")
        count: Number of symbols needed
        
    Returns:
        List of formatting symbols
    """
    if numbering_format == "i, ii, iii":
        roman_numerals = ["i", "ii", "iii", "iv", "v", "vi", "vii", "viii", "ix", "x", 
                         "xi", "xii", "xiii", "xiv", "xv", "xvi", "xvii", "xviii", "xix", "xx"]
        return roman_numerals[:count] + [str(i) for i in range(count - len(roman_numerals) + 1, count + 1) if i > len(roman_numerals)]
    
    elif numbering_format == "a, b, c":
        alphabet = ["a", "b", "c", "d", "e", "f", "g", "h", "i", "j",
                   "k", "l", "m", "n", "o", "p", "q", "r", "s", "t",
                   "u", "v", "w", "x", "y", "z"]
        return alphabet[:count] + [str(i) for i in range(count - len(alphabet) + 1, count + 1) if i > len(alphabet)]
    
    elif numbering_format == "*, †, ‡":
        symbols = ["*", "†", "‡", "§", "¶", "||", "**", "††", "‡‡", "§§"]
        return symbols[:count] + [str(i) for i in range(count - len(symbols) + 1, count + 1) if i > len(symbols)]
    
    else:  # Default to numbers
        return [str(i) for i in range(1, count + 1)]


def customize_footnote_formatting(doc, footnote_refs, format_symbols, start_number, style=None):
    """
    Apply custom formatting to footnote references and text.
    
    Args:
        doc: Document object
        footnote_refs: List of footnote references from find_footnote_references()
        format_symbols: List of formatting symbols to use
        start_number: Number to start footnote numbering from
        style: Optional style to apply to footnote text
        
    Returns:
        Number of footnotes formatted
    """
    # Update footnote references with new format
    for i, (para_idx, run_idx, _) in enumerate(footnote_refs):
        try:
            idx = i + start_number - 1
            if idx < len(format_symbols):
                symbol = format_symbols[idx]
            else:
                symbol = str(idx + 1)  # Fall back to numbers if we run out of symbols
            
            paragraph = doc.paragraphs[para_idx]
            paragraph.runs[run_idx].text = symbol
        except IndexError:
            # Skip if we can't locate the reference
            pass
    
    # Find footnote section and update
    found_footnote_section = False
    for para_idx, para in enumerate(doc.paragraphs):
        if para.text.startswith("Footnotes:") or para.text == "Footnotes":
            found_footnote_section = True
            
            # Update footnotes with new symbols
            for i in range(len(footnote_refs)):
                try:
                    footnote_para_idx = para_idx + i + 1
                    if footnote_para_idx < len(doc.paragraphs):
                        para = doc.paragraphs[footnote_para_idx]
                        
                        # Extract and preserve footnote text
                        footnote_text = para.text
                        if " " in footnote_text and len(footnote_text) > 2:
                            # Remove the old footnote number/symbol
                            footnote_text = footnote_text.split(" ", 1)[1]
                        
                        # Add new format
                        idx = i + start_number - 1
                        if idx < len(format_symbols):
                            symbol = format_symbols[idx]
                        else:
                            symbol = str(idx + 1)
                        
                        # Apply new formatting
                        para.text = f"{symbol} {footnote_text}"
                        
                        # Apply style
                        if style:
                            para.style = style
                except IndexError:
                    pass
            
            break
    
    return len(footnote_refs)
